from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = "docs/InternQuest_Artefact_SRS_Testing.docx"


sections = [
    {
        "number": "9.1",
        "title": "User Authentication and Access Control",
        "intro": (
            "The User Authentication and Access Control sub-system focuses on managing secure access to InternQuest. "
            "It allows users to register, log in, complete role-based onboarding, and access only the pages and "
            "features permitted for their role. This sub-system supports candidates, recruiters, and administrators "
            "by ensuring that each user is directed to the correct dashboard and protected from unauthorized access."
        ),
        "srs": [
            ("UAAC-F-1.0", "The system shall allow users to register using valid account details."),
            ("UAAC-F-1.1", "The system shall allow registered users to log in securely."),
            ("UAAC-F-1.2", "The system shall allow users to complete role-based onboarding."),
            ("UAAC-F-1.3", "The system shall assign users to candidate, recruiter, or administrator roles."),
            ("UAAC-F-1.4", "The system shall redirect users to dashboards based on their roles."),
            ("UAAC-F-1.5", "The system shall protect restricted routes from unauthorized users."),
            ("UAAC-F-1.6", "The system shall prevent suspended users from accessing protected features."),
            ("UAAC-NF-1.0", "Authentication responses shall be handled securely."),
            ("UAAC-NF-1.1", "Role-based access checks shall be performed consistently."),
            ("UAAC-NF-1.2", "The login and onboarding process shall be simple and responsive across devices."),
        ],
        "objectives": [
            "Ensure users can register and log in successfully.",
            "Verify role-based onboarding and dashboard redirection.",
            "Validate protected route and suspended account restrictions.",
        ],
        "tests": [
            ("UAAC-TC-1", "User registration", "Enter valid registration details", "Account created successfully"),
            ("UAAC-TC-2", "User login", "Enter valid login credentials", "User logged in successfully"),
            ("UAAC-TC-3", "Role-based onboarding", "Select user role and complete onboarding", "User role saved successfully"),
            ("UAAC-TC-4", "Dashboard redirection", "Log in as candidate/recruiter/admin", "Correct dashboard displayed"),
            ("UAAC-TC-5", "Route protection", "Access protected route without login", "User redirected to login page"),
            ("UAAC-TC-6", "Suspended account access", "Log in using suspended account", "Access restricted with suspension message"),
        ],
    },
    {
        "number": "9.2",
        "title": "Candidate Management Module",
        "intro": (
            "The Candidate Management Module focuses on enabling students to explore, apply for, and track internship "
            "opportunities. Candidates can build their profiles, upload resumes, search and filter opportunities, "
            "bookmark listings, submit applications, track application status, take assigned quizzes, and access "
            "learning resources. This module provides students with a structured way to manage their internship journey."
        ),
        "srs": [
            ("CMM-F-2.0", "The system shall allow candidates to build and update their profiles."),
            ("CMM-F-2.1", "The system shall allow candidates to search internship opportunities."),
            ("CMM-F-2.2", "The system shall allow candidates to filter opportunities by category and location type."),
            ("CMM-F-2.3", "The system shall allow candidates to bookmark internship listings."),
            ("CMM-F-2.4", "The system shall allow candidates to upload and manage resumes."),
            ("CMM-F-2.5", "The system shall allow candidates to submit internship applications."),
            ("CMM-F-2.6", "The system shall allow candidates to track application status."),
            ("CMM-F-2.7", "The system shall display assigned quizzes for selected candidates."),
            ("CMM-F-2.8", "The system shall allow candidates to access learning resources and sample quizzes."),
            ("CMM-NF-2.0", "Resume uploads shall be limited to supported file formats."),
            ("CMM-NF-2.1", "The candidate dashboard shall provide responsive performance across devices."),
            ("CMM-NF-2.2", "The system shall prevent duplicate applications for the same internship."),
        ],
        "objectives": [
            "Ensure candidates can search, apply, and track internships.",
            "Verify resume upload, bookmarking, and profile management.",
            "Validate quiz access and learning resource functionality.",
        ],
        "tests": [
            ("CMM-TC-1", "Build candidate profile", "Enter education, skills, experience, and preferences", "Profile saved successfully"),
            ("CMM-TC-2", "Search internship", "Enter keyword or category", "Relevant internships displayed"),
            ("CMM-TC-3", "Bookmark internship", "Click bookmark on internship listing", "Internship added to bookmarks"),
            ("CMM-TC-4", "Upload resume", "Upload supported resume file", "Resume uploaded successfully"),
            ("CMM-TC-5", "Apply for internship", "Submit application with resume and details", "Application submitted successfully"),
            ("CMM-TC-6", "Track application", "Open application history", "Correct application status displayed"),
            ("CMM-TC-7", "Attempt assigned quiz", "Start assigned quiz", "Quiz opens and accepts answers"),
            ("CMM-TC-8", "Access resource", "Open learning resource page", "Resource content displayed successfully"),
        ],
    },
    {
        "number": "9.3",
        "title": "Recruiter Management Module",
        "intro": (
            "The Recruiter Management Module allows recruiters to manage internship opportunities and evaluate "
            "applicants. Recruiters can create and manage internship listings, review candidate applications, update "
            "application statuses, create and assign quizzes, evaluate quiz results, and view recruitment analytics. "
            "This sub-system supports efficient candidate screening and internship management."
        ),
        "srs": [
            ("RMM-F-3.0", "The system shall allow recruiters to create internship listings."),
            ("RMM-F-3.1", "The system shall allow recruiters to edit and manage their own listings."),
            ("RMM-F-3.2", "The system shall allow recruiters to open or close internship listings."),
            ("RMM-F-3.3", "The system shall allow recruiters to view applications for their listings."),
            ("RMM-F-3.4", "The system shall allow recruiters to review candidate resumes and details."),
            ("RMM-F-3.5", "The system shall allow recruiters to update application statuses."),
            ("RMM-F-3.6", "The system shall allow recruiters to create and publish quizzes."),
            ("RMM-F-3.7", "The system shall allow recruiters to assign quizzes to applicants."),
            ("RMM-F-3.8", "The system shall allow recruiters to evaluate quiz results."),
            ("RMM-F-3.9", "The system shall display recruitment analytics to recruiters."),
            ("RMM-NF-3.0", "Recruiters shall only manage listings created by themselves."),
            ("RMM-NF-3.1", "Application review pages shall load efficiently."),
            ("RMM-NF-3.2", "Recruitment analytics shall be accurate and understandable."),
        ],
        "objectives": [
            "Ensure recruiters can create and manage internship listings.",
            "Verify application review and status update functionality.",
            "Validate quiz assignment, result evaluation, and analytics.",
        ],
        "tests": [
            ("RMM-TC-1", "Create internship", "Enter internship details and submit", "Internship listing created successfully"),
            ("RMM-TC-2", "Edit internship", "Modify existing listing details", "Updated details saved successfully"),
            ("RMM-TC-3", "Review application", "Open applicant details page", "Candidate details and resume displayed"),
            ("RMM-TC-4", "Update status", "Change application status", "New status saved and shown to candidate"),
            ("RMM-TC-5", "Create quiz", "Add quiz title, questions, and settings", "Quiz created successfully"),
            ("RMM-TC-6", "Assign quiz", "Assign quiz to an applicant", "Quiz assigned successfully"),
            ("RMM-TC-7", "Evaluate quiz result", "Open submitted quiz attempt", "Score and answers displayed"),
            ("RMM-TC-8", "View analytics", "Open recruiter analytics dashboard", "Recruitment metrics displayed correctly"),
        ],
    },
    {
        "number": "9.4",
        "title": "Administrative Management Module",
        "intro": (
            "The Administrative Management Module focuses on maintaining the quality, safety, and reliability of "
            "InternQuest. Administrators can manage users, suspend accounts, moderate internship listings, approve or "
            "reject categories, publish blog and resource content, create sample quizzes, and review reports. This "
            "module ensures that platform activities remain controlled and trustworthy."
        ),
        "srs": [
            ("AMM-F-4.0", "The system shall allow administrators to view and manage users."),
            ("AMM-F-4.1", "The system shall allow administrators to suspend user accounts."),
            ("AMM-F-4.2", "The system shall allow administrators to unsuspend user accounts."),
            ("AMM-F-4.3", "The system shall allow administrators to moderate internship listings."),
            ("AMM-F-4.4", "The system shall allow administrators to approve or reject internship categories."),
            ("AMM-F-4.5", "The system shall allow administrators to create and publish blog/resource posts."),
            ("AMM-F-4.6", "The system shall allow administrators to create sample quizzes."),
            ("AMM-F-4.7", "The system shall allow administrators to review submitted reports."),
            ("AMM-F-4.8", "The system shall allow administrators to take moderation actions on reported content."),
            ("AMM-NF-4.0", "Administrative features shall be accessible only to admin users."),
            ("AMM-NF-4.1", "Admin actions shall be stored reliably and reflected immediately."),
            ("AMM-NF-4.2", "Moderation pages shall be clear and easy to operate."),
        ],
        "objectives": [
            "Ensure administrators can manage users and content.",
            "Verify moderation of internships, categories, and reports.",
            "Validate admin-only access control.",
        ],
        "tests": [
            ("AMM-TC-1", "View users", "Open admin users page", "List of users displayed"),
            ("AMM-TC-2", "Suspend user", "Select user and submit suspension reason", "User account suspended successfully"),
            ("AMM-TC-3", "Unsuspend user", "Select suspended user and remove suspension", "User account restored successfully"),
            ("AMM-TC-4", "Moderate internship", "Close inappropriate internship listing", "Listing closed successfully"),
            ("AMM-TC-5", "Review category", "Approve or reject requested category", "Category status updated"),
            ("AMM-TC-6", "Publish resource", "Create and publish blog/resource post", "Resource available to users"),
            ("AMM-TC-7", "Review report", "Open report and submit decision", "Report status updated successfully"),
            ("AMM-TC-8", "Admin route access", "Access admin page as non-admin user", "Access denied or redirected"),
        ],
    },
    {
        "number": "9.5",
        "title": "Shared Platform Services",
        "intro": (
            "The Shared Platform Services sub-system provides common services used across InternQuest. It includes "
            "in-app notifications, email notifications, dashboard views, reporting and analytics, file storage, and "
            "content/data management. These services support communication, monitoring, and data handling for "
            "candidates, recruiters, and administrators."
        ),
        "srs": [
            ("SPS-F-5.0", "The system shall generate in-app notifications for important user events."),
            ("SPS-F-5.1", "The system shall send email notifications for application, quiz, resource, and internship updates."),
            ("SPS-F-5.2", "The system shall display dashboard summaries based on user role."),
            ("SPS-F-5.3", "The system shall track internship views and application activity."),
            ("SPS-F-5.4", "The system shall generate recruitment and platform analytics."),
            ("SPS-F-5.5", "The system shall support secure file upload and retrieval."),
            ("SPS-F-5.6", "The system shall manage shared content and data across modules."),
            ("SPS-NF-5.0", "Notifications shall be delivered reliably."),
            ("SPS-NF-5.1", "Uploaded files shall be securely stored and accessed."),
            ("SPS-NF-5.2", "Dashboard and analytics data shall be updated accurately."),
            ("SPS-NF-5.3", "Shared services shall support all user roles consistently."),
        ],
        "objectives": [
            "Ensure notifications and emails are triggered correctly.",
            "Verify dashboard, analytics, and file storage functionality.",
            "Validate shared services across different user roles.",
        ],
        "tests": [
            ("SPS-TC-1", "In-app notification", "Trigger application or quiz event", "Notification appears for target user"),
            ("SPS-TC-2", "Email notification", "Trigger email-supported event", "Email notification sent successfully"),
            ("SPS-TC-3", "Dashboard summary", "Open dashboard as candidate/recruiter/admin", "Correct role-based summary displayed"),
            ("SPS-TC-4", "Track internship view", "Open internship detail page", "View count updated correctly"),
            ("SPS-TC-5", "View analytics", "Open analytics dashboard", "Correct metrics displayed"),
            ("SPS-TC-6", "Upload file", "Upload supported document or image", "File uploaded successfully"),
            ("SPS-TC-7", "Retrieve file", "Open uploaded file link", "File displayed or downloaded correctly"),
            ("SPS-TC-8", "Shared service access", "Use service from different roles", "Service works according to role permission"),
        ],
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def format_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, sum(widths))
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], "F2F2F2")
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = value

    format_table(table, widths)
    document.add_paragraph()


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.35)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.color.rgb = None
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Artefact Designs: SRS and Testing")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(16)

    for index, item in enumerate(sections):
        if index > 0:
            doc.add_section(WD_SECTION.NEW_PAGE)

        doc.add_heading(f"{item['number']} {item['title']}", level=1)
        p = doc.add_paragraph(item["intro"])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_heading(f"{item['number']}.1 SRS", level=2)
        add_table(
            doc,
            ["Req. Code", "Req. Desc"],
            item["srs"],
            [1500, 7860],
        )

        doc.add_heading(f"{item['number']}.3 Test Cases", level=2)
        doc.add_paragraph("Test Objectives")
        for objective in item["objectives"]:
            add_bullet(doc, objective)

        add_table(
            doc,
            ["Test Case ID", "Test Scenario", "Test Steps", "Expected Result"],
            item["tests"],
            [1400, 2300, 2800, 2860],
        )

    doc.save(OUT)


if __name__ == "__main__":
    main()
